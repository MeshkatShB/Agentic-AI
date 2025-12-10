"""Utility functions for parsing uploaded files."""

import io
from pathlib import Path
from typing import Dict, Any, Tuple, Optional
import pypdf
import docx
import markdown
from bs4 import BeautifulSoup
import logging

logger = logging.getLogger(__name__)


async def extract_file_content(
    file_content: bytes,
    filename: str,
    extract_metadata: bool = False
) -> Tuple[str, Dict[str, Any]]:
    """Extract text content from uploaded file.
    
    Args:
        file_content: File content as bytes
        filename: Original filename
        extract_metadata: Whether to extract metadata
        
    Returns:
        Tuple of (content, metadata)
    """
    ext = Path(filename).suffix.lower()
    content = ""
    metadata = {
        "filename": filename,
        "file_type": ext,
        "file_size": len(file_content)
    }
    
    try:
        if ext == ".pdf":
            content, pdf_metadata = _parse_pdf_bytes(file_content, extract_metadata)
            if extract_metadata:
                metadata.update(pdf_metadata)
        elif ext == ".docx":
            content, docx_metadata = _parse_docx_bytes(file_content, extract_metadata)
            if extract_metadata:
                metadata.update(docx_metadata)
        elif ext == ".md":
            content, md_metadata = _parse_markdown_bytes(file_content, extract_metadata)
            if extract_metadata:
                metadata.update(md_metadata)
        elif ext in [".txt", ".log", ".csv", ".json", ".xml", ".html", ".htm"]:
            content = file_content.decode('utf-8', errors='ignore')
            if extract_metadata:
                metadata.update({
                    "lines": len(content.splitlines()),
                    "characters": len(content)
                })
        else:
            # Try to decode as text for unknown types
            try:
                content = file_content.decode('utf-8', errors='ignore')
                if extract_metadata:
                    metadata.update({
                        "lines": len(content.splitlines()),
                        "characters": len(content)
                    })
            except Exception:
                raise ValueError(f"Unsupported file type: {ext}")
        
        return content, metadata
        
    except Exception as e:
        logger.error(f"Failed to extract content from {filename}: {e}")
        raise


def _parse_pdf_bytes(file_content: bytes, extract_metadata: bool) -> Tuple[str, Dict[str, Any]]:
    """Parse PDF from bytes."""
    content = []
    metadata = {}
    
    file_stream = io.BytesIO(file_content)
    reader = pypdf.PdfReader(file_stream)
    
    if extract_metadata:
        info = reader.metadata
        if info:
            metadata = {
                "title": info.get('/Title', ''),
                "author": info.get('/Author', ''),
                "subject": info.get('/Subject', ''),
                "pages": len(reader.pages)
            }
    
    for page in reader.pages:
        content.append(page.extract_text())
    
    return '\n'.join(content), metadata


def _parse_docx_bytes(file_content: bytes, extract_metadata: bool) -> Tuple[str, Dict[str, Any]]:
    """Parse DOCX from bytes."""
    file_stream = io.BytesIO(file_content)
    doc = docx.Document(file_stream)
    content = []
    metadata = {}
    
    for para in doc.paragraphs:
        content.append(para.text)
    
    if extract_metadata:
        props = doc.core_properties
        metadata = {
            "title": props.title or '',
            "author": props.author or '',
            "created": props.created.isoformat() if props.created else '',
            "modified": props.modified.isoformat() if props.modified else '',
            "paragraphs": len(doc.paragraphs)
        }
    
    return '\n'.join(content), metadata


def _parse_markdown_bytes(file_content: bytes, extract_metadata: bool) -> Tuple[str, Dict[str, Any]]:
    """Parse Markdown from bytes."""
    text = file_content.decode('utf-8', errors='ignore')
    html = markdown.markdown(text)
    soup = BeautifulSoup(html, 'html.parser')
    content = soup.get_text()
    
    metadata = {}
    if extract_metadata:
        metadata = {
            "headings": len(soup.find_all(['h1', 'h2', 'h3'])),
            "links": len(soup.find_all('a')),
            "code_blocks": len(soup.find_all('code'))
        }
    
    return content, metadata

