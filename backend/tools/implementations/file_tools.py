"""File-related tools."""

import os
import re
import hashlib
from pathlib import Path
from typing import Dict, Any, List
from datetime import datetime
import pypdf
import docx
import markdown
from bs4 import BeautifulSoup
from backend.tools.base import BaseTool, ToolPermission, ToolResult
from backend.config import settings
from backend.storage import get_vector_store
import logging

logger = logging.getLogger(__name__)


class SearchLocalFilesTool(BaseTool):
    """Simple RAG tool for retrieving information from indexed local files.
    
    Follows LangChain RAG agent pattern: uses semantic search to retrieve relevant
    documents from the vector store and returns them in a format suitable for LLM context.
    """
    
    def __init__(self):
        super().__init__()
        self.vector_store = None
        
    @property
    def name(self) -> str:
        return "search_local_files"
    
    @property
    def description(self) -> str:
        return "Retrieve information related to a query by searching through indexed local files. Use this tool to find relevant documents and content that can help answer user questions."
    
    @property
    def parameters(self) -> Dict[str, Any]:
        return {
            "type": "object",
            "properties": {
                "query": {
                    "type": "string",
                    "description": "The search query or question to find relevant information"
                },
                "user_id": {
                    "type": "integer",
                    "description": "User ID for searching user-specific documents (optional, defaults to searching user documents if available)"
                }
            },
            "required": ["query"]
        }
    
    @property
    def permission(self) -> ToolPermission:
        return ToolPermission.READ_FILES
    
    async def _get_vector_store(self):
        """Get or initialize vector store."""
        if self.vector_store is None:
            self.vector_store = get_vector_store()
        return self.vector_store
    
    async def execute(self, **kwargs) -> ToolResult:
        """Execute RAG retrieval following LangChain pattern.
        
        Retrieves relevant documents from the vector store using semantic search
        and returns them in a serialized format suitable for LLM context.
        Searches user-specific document collections if user_id is provided.
        """
        query = kwargs.get("query", "").strip()
        user_id = kwargs.get("user_id")
        
        if not query:
            return ToolResult(
                success=False,
                output=None,
                error="Query cannot be empty"
            )
        
        try:
            # Get vector store and ensure it's initialized
            vector_store = await self._get_vector_store()
            await vector_store.initialize()
            
            # Get user's embedding model preference if user_id is provided
            embedding_model = None
            if user_id:
                from backend.models import SessionLocal, User
                db = SessionLocal()
                try:
                    user = db.query(User).filter(User.id == user_id).first()
                    if user and user.preferences:
                        embedding_model = user.preferences.get("embedding_model")
                finally:
                    db.close()
            
            # Determine collection name - prefer user-specific documents
            collection_name = "documents"  # Default to global collection
            if user_id:
                collection_name = f"user_documents_{user_id}"
                # Ensure user collection exists (it will be created on first document upload)
                # Try to search user documents
                try:
                    retrieved_docs = await vector_store.search(
                        query=query,
                        k=5,
                        collection_name=collection_name,
                        embedding_model=embedding_model
                    )
                except Exception as e:
                    # Collection might not exist yet if user has no documents
                    logger.debug(f"User collection {collection_name} not found or empty: {e}")
                    retrieved_docs = []
                
                # If no user documents found, return empty result
                # User wants to search ONLY their documents, not the global ./data directory
                if not retrieved_docs:
                    return ToolResult(
                        success=True,
                        output="No documents found matching the query in your uploaded documents. Please upload documents in the Documents tab to enable search.",
                        metadata={"query": query, "results_count": 0, "collection": collection_name}
                    )
            else:
                # No user_id provided, fall back to global collection (legacy behavior)
                retrieved_docs = await vector_store.search(
                    query=query,
                    k=5,
                    collection_name="documents",
                    embedding_model=embedding_model
                )
            
            if not retrieved_docs:
                return ToolResult(
                    success=True,
                    output="No documents found matching the query.",
                    metadata={"query": query, "results_count": 0}
                )
            
            # Serialize documents in LangChain format: "Source: {source}\nContent: {content}"
            serialized_docs = []
            for doc in retrieved_docs:
                source = doc.get("metadata", {}).get("file_name", "Unknown")
                content = doc.get("document", "")
                serialized = f"Source: {source}\nContent: {content}"
                serialized_docs.append(serialized)
            
            # Join all serialized documents
            serialized_output = "\n\n".join(serialized_docs)
            
            logger.info(f"Retrieved {len(retrieved_docs)} documents for query: '{query}'")
            
            return ToolResult(
                success=True,
                output=serialized_output,
                metadata={
                    "query": query,
                    "results_count": len(retrieved_docs),
                    "documents": retrieved_docs  # Include raw documents for reference
                }
            )
            
        except Exception as e:
            logger.error(f"RAG retrieval failed: {e}", exc_info=True)
            return ToolResult(
                success=False,
                output=None,
                error=f"Search failed: {str(e)}"
            )


class ReadFileTool(BaseTool):
    """Tool for reading file contents."""
    
    @property
    def name(self) -> str:
        return "read_file"
    
    @property
    def description(self) -> str:
        return "Read the contents of a specific file"
    
    @property
    def parameters(self) -> Dict[str, Any]:
        return {
            "type": "object",
            "properties": {
                "file_path": {
                    "type": "string",
                    "description": "Path to the file to read"
                },
                "start_line": {
                    "type": "integer",
                    "description": "Start line number (for text files)",
                    "default": 1
                },
                "end_line": {
                    "type": "integer",
                    "description": "End line number (for text files)",
                    "default": -1
                }
            },
            "required": ["file_path"]
        }
    
    @property
    def permission(self) -> ToolPermission:
        return ToolPermission.READ_FILES
    
    async def execute(self, **kwargs) -> ToolResult:
        """Read the file."""
        file_path = kwargs.get("file_path")
        start_line = kwargs.get("start_line", 1)
        end_line = kwargs.get("end_line", -1)
        
        try:
            path = Path(file_path).resolve()
            
            # Check if path is allowed
            if not settings.is_path_allowed(path):
                return ToolResult(
                    success=False,
                    output=None,
                    error=f"Access denied: {file_path}"
                )
            
            # Check if file exists
            if not path.exists():
                return ToolResult(
                    success=False,
                    output=None,
                    error=f"File not found: {file_path}"
                )
            
            # Read file
            with open(path, 'r', encoding='utf-8') as f:
                lines = f.readlines()
            
            # Apply line filtering
            if end_line == -1:
                content = ''.join(lines[start_line-1:])
            else:
                content = ''.join(lines[start_line-1:end_line])
            
            return ToolResult(
                success=True,
                output=content,
                metadata={
                    "file_path": str(path),
                    "total_lines": len(lines),
                    "size_bytes": path.stat().st_size
                }
            )
            
        except Exception as e:
            logger.error(f"Read file failed: {e}")
            return ToolResult(
                success=False,
                output=None,
                error=str(e)
            )


class ParseDocumentTool(BaseTool):
    """Tool for parsing various document formats."""
    
    @property
    def name(self) -> str:
        return "parse_document"
    
    @property
    def description(self) -> str:
        return "Extract text from PDF, DOCX, or other document formats"
    
    @property
    def parameters(self) -> Dict[str, Any]:
        return {
            "type": "object",
            "properties": {
                "file_path": {
                    "type": "string",
                    "description": "Path to the document to parse"
                },
                "extract_metadata": {
                    "type": "boolean",
                    "description": "Extract document metadata",
                    "default": False
                }
            },
            "required": ["file_path"]
        }
    
    @property
    def permission(self) -> ToolPermission:
        return ToolPermission.READ_FILES
    
    async def execute(self, **kwargs) -> ToolResult:
        """Parse the document."""
        file_path = kwargs.get("file_path")
        extract_metadata = kwargs.get("extract_metadata", False)
        
        try:
            path = Path(file_path).resolve()
            
            # Check if path is allowed
            if not settings.is_path_allowed(path):
                return ToolResult(
                    success=False,
                    output=None,
                    error=f"Access denied: {file_path}"
                )
            
            # Check if file exists
            if not path.exists():
                return ToolResult(
                    success=False,
                    output=None,
                    error=f"File not found: {file_path}"
                )
            
            # Determine file type and parse
            ext = path.suffix.lower()
            content = ""
            metadata = {}
            
            if ext == ".pdf":
                content, metadata = self._parse_pdf(path, extract_metadata)
            elif ext == ".docx":
                content, metadata = self._parse_docx(path, extract_metadata)
            elif ext == ".md":
                content, metadata = self._parse_markdown(path, extract_metadata)
            elif ext in [".txt", ".log", ".csv"]:
                content = path.read_text(encoding='utf-8')
                if extract_metadata:
                    metadata = {
                        "lines": len(content.splitlines()),
                        "characters": len(content)
                    }
            else:
                return ToolResult(
                    success=False,
                    output=None,
                    error=f"Unsupported file type: {ext}"
                )
            
            return ToolResult(
                success=True,
                output={
                    "content": content,
                    "metadata": metadata if extract_metadata else {}
                },
                metadata={
                    "file_type": ext,
                    "file_size": path.stat().st_size
                }
            )
            
        except Exception as e:
            logger.error(f"Parse document failed: {e}")
            return ToolResult(
                success=False,
                output=None,
                error=str(e)
            )
    
    def _parse_pdf(self, path: Path, extract_metadata: bool):
        """Parse PDF file."""
        content = []
        metadata = {}
        
        with open(path, 'rb') as f:
            reader = pypdf.PdfReader(f)
            
            if extract_metadata:
                info = reader.metadata
                if info:
                    metadata = {
                        "title": info.get('/Title', ''),
                        "author": info.get('/Author', ''),
                        "subject": info.get('/Subject', ''),
                        "pages": len(reader.pages)
                    }
            
            for page in reader.pages:
                content.append(page.extract_text())
        
        return '\n'.join(content), metadata
    
    def _parse_docx(self, path: Path, extract_metadata: bool):
        """Parse DOCX file."""
        doc = docx.Document(path)
        content = []
        metadata = {}
        
        for para in doc.paragraphs:
            content.append(para.text)
        
        if extract_metadata:
            props = doc.core_properties
            metadata = {
                "title": props.title or '',
                "author": props.author or '',
                "created": props.created.isoformat() if props.created else '',
                "modified": props.modified.isoformat() if props.modified else '',
                "paragraphs": len(doc.paragraphs)
            }
        
        return '\n'.join(content), metadata
    
    def _parse_markdown(self, path: Path, extract_metadata: bool):
        """Parse Markdown file."""
        content = path.read_text(encoding='utf-8')
        html = markdown.markdown(content)
        soup = BeautifulSoup(html, 'html.parser')
        text = soup.get_text()
        
        metadata = {}
        if extract_metadata:
            metadata = {
                "headings": len(soup.find_all(['h1', 'h2', 'h3'])),
                "links": len(soup.find_all('a')),
                "code_blocks": len(soup.find_all('code'))
            }
        
        return text, metadata
